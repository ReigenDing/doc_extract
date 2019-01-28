import subprocess
import chardet

import docx
file=docx.Document('D:\work\搜索demo-文档\原始数据\数据底座数据主题联接元数据注册指南(试行)-通知【2017】001.docx')
#获取文档段落数
print(len(file.paragraphs))
#输出每一段的内容
with open('D:\work\文档解析\data\input\数据底座数据主题联接元数据注册指南(试行)-通知【2017】001.txt', 'w', encoding='utf8') as fp:
    for para in file.paragraphs:
        print(para.text)
        fp.write(para.text + '\n')




